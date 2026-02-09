"""
Formatting utilities for DocFlow.
Handles stats counting, quality scoring, and format conversions.
"""

import os
import json
import markdown2
from docx import Document
from datetime import datetime

# Constants
STATS_FILE = "usage_stats.json"


def count_stats(markdown_text):
    """Count words and characters in markdown text."""
    words = len(markdown_text.split())
    chars = len(markdown_text)
    return words, chars


def estimate_quality_score(markdown_text, method_used):
    """Estimate quality score based on content and method."""
    score = 50  # Base score
    
    # Method bonus
    if "OpenRouter" in method_used:
        score += 30
    elif "GMFT" in method_used:
        score += 25
    elif "RapidOCR" in method_used:
        score += 20
    elif "PyMuPDF" in method_used:
        score += 15
    
    # Content quality checks
    if "##" in markdown_text:  # Has headings
        score += 10
    if "|" in markdown_text:  # Has tables
        score += 10
    if len(markdown_text) > 1000:  # Substantial content
        score += 10
    
    return min(100, score)


def estimate_cost(model, num_pages):
    """Estimate cost for OpenRouter models."""
    cost_per_1k = {
        "Nemotron Nano 12B VL (FREE)": 0,
        "Gemini 2.0 Flash Lite ($0.08/1K pages)": 0.08,
        "Qwen 2.5-VL 32B ($0.05/1K pages)": 0.05,
        "Qwen 2.5-VL 72B ($0.15/1K pages)": 0.15,
        "Mistral Pixtral Large ($2/1K pages)": 2.0
    }
    
    rate = cost_per_1k.get(model, 0)
    cost = (num_pages / 1000) * rate
    
    if cost == 0:
        return "**Estimated Cost:** FREE"
    elif cost < 0.01:
        return f"**Estimated Cost:** < $0.01"
    else:
        return f"**Estimated Cost:** ${cost:.2f}"


def markdown_to_html(markdown_text):
    """Convert markdown to HTML."""
    return markdown2.markdown(markdown_text, extras=["tables", "fenced-code-blocks"])


def markdown_to_txt(markdown_text):
    """Convert markdown to plain text."""
    # Simple conversion - remove markdown syntax
    import re
    text = markdown_text
    text = re.sub(r'#{1,6}\s', '', text)  # Remove headers
    text = re.sub(r'\*\*(.+?)\*\*', r'\1', text)  # Remove bold
    text = re.sub(r'\*(.+?)\*', r'\1', text)  # Remove italic
    text = re.sub(r'\[(.+?)\]\(.+?\)', r'\1', text)  # Remove links
    return text


def markdown_to_docx(markdown_text, output_path):
    """Convert markdown to Word document."""
    doc = Document()
    lines = markdown_text.split('\n')
    
    for line in lines:
        if line.startswith('# '):
            doc.add_heading(line[2:], level=1)
        elif line.startswith('## '):
            doc.add_heading(line[3:], level=2)
        elif line.startswith('### '):
            doc.add_heading(line[4:], level=3)
        elif line.strip():
            doc.add_paragraph(line)
    
    doc.save(output_path)


def load_stats():
    """Load usage statistics from file."""
    if os.path.exists(STATS_FILE):
        try:
            with open(STATS_FILE, 'r', encoding='utf-8') as f:
                return json.load(f)
        except (json.JSONDecodeError, IOError, OSError) as e:
            # Stats file is corrupted or unreadable - log and return defaults
            print(f"Warning: Failed to load stats file {STATS_FILE}: {type(e).__name__} - {e}")
            pass
    return {"total_conversions": 0, "total_words": 0, "total_files": 0, "avg_time": 0}


def save_stats(stats):
    """Save usage statistics to file."""
    with open(STATS_FILE, 'w', encoding='utf-8') as f:
        json.dump(stats, f, indent=2)


def update_stats(words, elapsed_time):
    """Update usage statistics with new conversion."""
    stats = load_stats()
    stats["total_conversions"] += 1
    stats["total_words"] += words
    stats["total_files"] += 1
    n = stats["total_conversions"]
    stats["avg_time"] = ((stats["avg_time"] * (n-1)) + elapsed_time) / n
    save_stats(stats)
    return stats


def get_stats_display():
    """Generate formatted stats display for UI."""
    stats = load_stats()
    return f"""**Usage Statistics**

**Total Conversions:** {stats['total_conversions']}
**Total Words:** {stats['total_words']:,}
**Avg Time:** {stats['avg_time']:.1f}s
"""
